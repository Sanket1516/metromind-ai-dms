"""
MetroMind Document Service
Document upload, processing, and management
"""
from sqlalchemy import text
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form, Request, status, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel, field_validator
from sqlalchemy import or_
from sqlalchemy.orm import Session
from datetime import datetime, timezone
from typing import Optional, List, Dict, Any, Union
import os
import shutil
import hashlib
import uuid
import mimetypes
import asyncio
from pathlib import Path
import json

# Import our models and database
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from database import (
    get_db, Document, User, AuditLog, Notification, DocumentStatus, 
    DocumentCategory, Priority, NotificationType, Permission, UserRole,SharedDocument,DocumentVersion
)
from config import service_config, app_config, ai_config
from utils.logging_utils import setup_logger
from services.auth_service import (
    get_current_user, require_permission, has_permission
)

# Setup
logger = setup_logger(__name__)
app = FastAPI(
    title="MetroMind Document Service",
    description="Document upload, processing, and management",
    version="1.0.0"
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    file_size: int
    status: str
    message: str

class DocumentInfo(BaseModel):
    id: str
    filename: str
    original_filename: str
    file_size: int
    mime_type: str
    title: Optional[str]
    description: Optional[str]
    category: str
    priority: str
    status: str
    uploaded_by: str
    created_at: datetime
    processed_at: Optional[datetime]
    extracted_text: Optional[str]
    summary: Optional[str]
    language_detected: Optional[str]
    tags: List[str] = []

class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    category: Optional[DocumentCategory] = None
    priority: Optional[Priority] = None
    tags: Optional[List[str]] = None

class DocumentSearch(BaseModel):
    query: Optional[str] = None
    category: Optional[DocumentCategory] = None
    priority: Optional[Priority] = None
    status: Optional[DocumentStatus] = None
    uploaded_by: Optional[str] = None
    date_from: Optional[datetime] = None
    date_to: Optional[datetime] = None
    limit: int = 20
    offset: int = 0
    include_shared: bool = True
    shared_by_me: bool = False
    shared_with_me: bool = False
    shared_with_department: bool = False

class ShareDocumentRequest(BaseModel):
    shared_with_user: Optional[str] = None  # User ID
    shared_with_department: Optional[str] = None  # Department name
    can_edit: bool = False
    expires_at: Optional[datetime] = None

    @field_validator('shared_with_department')
    @classmethod
    def validate_department(cls, v, info):
        values = info.data if hasattr(info, 'data') else {}
        if v is None and values.get('shared_with_user') is None:
            raise ValueError("Either user or department must be specified")
        if v is not None and values.get('shared_with_user') is not None:
            raise ValueError("Cannot specify both user and department")
        return v

class SharedDocumentInfo(BaseModel):
    id: str
    document_id: str
    document_title: Optional[str]
    document_filename: str
    shared_by: str
    shared_by_name: str
    shared_with_user: Optional[str]
    shared_with_user_name: Optional[str]
    shared_with_department: Optional[str]
    can_edit: bool
    created_at: datetime
    expires_at: Optional[datetime]

class SharedDocumentsList(BaseModel):
    shared_by_me: List[SharedDocumentInfo]
    shared_with_me: List[SharedDocumentInfo]

class DocumentVersionInfo(BaseModel):
    id: str
    document_id: str
    version_number: int
    file_size: int
    modified_by: str
    modified_by_name: str
    changes_description: Optional[str]
    created_at: datetime
    
class DocumentVersionList(BaseModel):
    versions: List[DocumentVersionInfo]
    
# Configuration
ALLOWED_EXTENSIONS = {
    '.pdf', '.doc', '.docx', '.txt', '.jpg', '.jpeg', '.png', '.bmp', '.tiff',
    '.xlsx', '.xls', '.csv', '.ppt', '.pptx', '.odt', '.rtf'
}

MAX_FILE_SIZE = ai_config.max_document_size_mb * 1024 * 1024  # Convert to bytes

# Utility functions
def get_file_hash(file_path: str) -> str:
    """Calculate SHA-256 hash of file"""
    hash_sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_sha256.update(chunk)
    return hash_sha256.hexdigest()

def validate_file(file: UploadFile) -> tuple[bool, str]:
    """Validate uploaded file"""
    # Check file extension
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        return False, f"File type {file_ext} not allowed"
    
    # Check file size (approximate)
    if hasattr(file, 'size') and file.size > MAX_FILE_SIZE:
        return False, f"File size exceeds maximum limit of {ai_config.max_document_size_mb}MB"
    
    # Check filename
    if not file.filename or len(file.filename) > 255:
        return False, "Invalid filename"
    
    return True, "Valid"

def determine_category(filename: str, content: str = "") -> DocumentCategory:
    """Determine document category based on filename and content"""
    filename_lower = filename.lower()
    content_lower = content.lower()
    
    # Safety keywords
    safety_keywords = [
        'safety', 'accident', 'incident', 'hazard', 'emergency', 'evacuation',
        'fire', 'security', 'risk', 'unsafe', 'danger'
    ]
    
    # Maintenance keywords
    maintenance_keywords = [
        'maintenance', 'repair', 'inspection', 'service', 'equipment', 'technical',
        'mechanical', 'electrical', 'fault', 'breakdown', 'preventive'
    ]
    
    # Finance keywords
    finance_keywords = [
        'budget', 'finance', 'payment', 'invoice', 'cost', 'expense', 'revenue',
        'accounting', 'audit', 'financial', 'procurement', 'purchase'
    ]
    
    # Operations keywords
    operations_keywords = [
        'operation', 'schedule', 'timetable', 'passenger', 'service', 'performance',
        'capacity', 'frequency', 'route', 'station', 'platform'
    ]
    
    # HR keywords
    hr_keywords = [
        'hr', 'human', 'employee', 'staff', 'training', 'recruitment', 'salary',
        'leave', 'policy', 'personnel', 'attendance'
    ]
    
    # Legal keywords
    legal_keywords = [
        'legal', 'contract', 'agreement', 'compliance', 'regulation', 'law',
        'litigation', 'terms', 'conditions', 'policy'
    ]
    
    # Regulatory keywords
    regulatory_keywords = [
        'regulatory', 'commissioner', 'ministry', 'government', 'compliance',
        'standard', 'specification', 'guideline', 'directive'
    ]
    
    text_to_check = f"{filename_lower} {content_lower}"
    
    # Check against keywords
    if any(keyword in text_to_check for keyword in safety_keywords):
        return DocumentCategory.SAFETY
    elif any(keyword in text_to_check for keyword in maintenance_keywords):
        return DocumentCategory.MAINTENANCE
    elif any(keyword in text_to_check for keyword in finance_keywords):
        return DocumentCategory.FINANCE
    elif any(keyword in text_to_check for keyword in operations_keywords):
        return DocumentCategory.OPERATIONS
    elif any(keyword in text_to_check for keyword in hr_keywords):
        return DocumentCategory.HR
    elif any(keyword in text_to_check for keyword in legal_keywords):
        return DocumentCategory.LEGAL
    elif any(keyword in text_to_check for keyword in regulatory_keywords):
        return DocumentCategory.REGULATORY
    
    return DocumentCategory.OTHER

def determine_priority(filename: str, content: str = "") -> Priority:
    """Determine document priority based on filename and content"""
    filename_lower = filename.lower()
    content_lower = content.lower()
    
    critical_keywords = ['emergency', 'urgent', 'critical', 'immediate', 'asap', 'crisis']
    high_keywords = ['important', 'priority', 'high', 'attention', 'escalate']
    
    text_to_check = f"{filename_lower} {content_lower}"
    
    if any(keyword in text_to_check for keyword in critical_keywords):
        return Priority.CRITICAL
    elif any(keyword in text_to_check for keyword in high_keywords):
        return Priority.HIGH
    
    return Priority.MEDIUM

def check_document_access(document: Document, user: User, db: Session, require_edit: bool = False) -> bool:
    """Check if user has access to document"""
    # Admin has full access
    if user.role == UserRole.ADMIN:
        return True
        
    # Document owner has full access
    if document.uploaded_by == user.id:
        return True
        
    # Check sharing permissions
    share = db.query(SharedDocument).filter(
        SharedDocument.document_id == document.id,
        (
            (SharedDocument.shared_with_user == user.id) |
            (SharedDocument.shared_with_department == user.department)
        ),
        (SharedDocument.expires_at.is_(None) | (SharedDocument.expires_at > datetime.now(timezone.utc)))
    ).first()
    
    if share:
        return not require_edit or share.can_edit
    
    # Check department and role based access
    if document.access_level == 1:  # Public
        return True
    elif document.access_level == 2:  # Department
        return document.uploaded_by_user.department == user.department
    elif document.access_level == 3:  # Manager
        return user.role in [UserRole.MANAGER, UserRole.SUPERVISOR]
    elif document.access_level == 4:  # Admin only
        return False  # Already checked admin above
        
    return False

async def save_uploaded_file(file: UploadFile, user_id: str, version_dir: Optional[str] = None) -> tuple[str, str]:
    """Save uploaded file to disk and return file path and hash"""
    # Create user-specific directory
    base_dir = Path(app_config.upload_directory)
    if version_dir:
        user_dir = base_dir / version_dir
    else:
        user_dir = base_dir / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate unique filename
    file_ext = Path(file.filename).suffix
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = user_dir / unique_filename
    
    # Save file
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # Calculate file hash
    file_hash = get_file_hash(str(file_path))
    
    return str(file_path), file_hash

async def process_document_content(document: Document, db: Session):
    """Process document content (OCR, text extraction, etc.)"""
    try:
        # Update status to processing
        document.status = DocumentStatus.PROCESSING
        document.processing_progress = 10
        db.commit()
        
        # Basic text extraction based on file type
        extracted_text = ""
        
        if document.mime_type == 'application/pdf':
            extracted_text = await extract_pdf_text(document.file_path)
        elif document.mime_type.startswith('image/'):
            extracted_text = await extract_image_text(document.file_path)
        elif document.mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                                    'application/msword']:
            extracted_text = await extract_word_text(document.file_path)
        elif document.mime_type == 'text/plain':
            with open(document.file_path, 'r', encoding='utf-8') as f:
                extracted_text = f.read()
        
        document.processing_progress = 50
        document.extracted_text = extracted_text[:10000]  # Limit to 10KB
        db.commit()
        
        # Auto-categorize if not already set
        if not document.category or document.category == DocumentCategory.OTHER:
            document.category = determine_category(document.original_filename, extracted_text)
        
        # Auto-prioritize if not already set
        if document.priority == Priority.MEDIUM:
            document.priority = determine_priority(document.original_filename, extracted_text)
        
        document.processing_progress = 80
        db.commit()
        
        # Generate basic summary
        if extracted_text:
            summary = generate_simple_summary(extracted_text)
            document.summary = summary
        
        # Mark as processed
        document.status = DocumentStatus.PROCESSED
        document.processing_progress = 100
        document.processed_at = datetime.now(timezone.utc)
        db.commit()
        
        logger.info(f"Document {document.id} processed successfully")
        
        # Send notifications if high priority
        if document.priority in [Priority.HIGH, Priority.CRITICAL]:
            await send_document_notification(document, db)
        
    except Exception as e:
        logger.error(f"Error processing document {document.id}: {e}")
        document.status = DocumentStatus.FAILED
        document.processing_error = str(e)
        db.commit()

async def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages[:5]:  # First 5 pages only
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return ""

async def extract_image_text(file_path: str) -> str:
    """Extract text from image using OCR (placeholder)"""
    try:
        # This would call the OCR service in a real implementation
        # For now, return placeholder
        return f"[OCR processing needed for {file_path}]"
    except Exception as e:
        logger.error(f"Error extracting image text: {e}")
        return ""

async def extract_word_text(file_path: str) -> str:
    """Extract text from Word document"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs[:50]:  # First 50 paragraphs only
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting Word text: {e}")
        return ""

def generate_simple_summary(text: str, max_length: int = 500) -> str:
    """Generate simple summary by taking first sentences"""
    if len(text) <= max_length:
        return text
    
    # Split into sentences and take first few
    sentences = text.split('.')[:3]
    summary = '. '.join(sentences)
    
    if len(summary) > max_length:
        summary = summary[:max_length] + "..."
    
    return summary

async def send_document_notification(document: Document, db: Session):
    """Send notification for high-priority documents"""
    try:
        # Get users who should be notified based on category and priority
        from database import UserRole
        
        # For high/critical priority, notify managers and admins
        target_roles = [UserRole.ADMIN, UserRole.MANAGER]
        
        users = db.query(User).filter(
            User.role.in_(target_roles),
            User.status == 'active'
        ).all()
        
        for user in users:
            notification = Notification(
                user_id=user.id,
                document_id=document.id,
                title=f"High Priority Document: {document.original_filename}",
                message=f"A {document.priority.value} priority document in {document.category.value} category has been uploaded.",
                notification_type=NotificationType.WARNING if document.priority == Priority.HIGH else NotificationType.ERROR
            )
            db.add(notification)
        
        db.commit()
        logger.info(f"Notifications sent for document {document.id}")
        
    except Exception as e:
        logger.error(f"Error sending document notification: {e}")

# API Endpoints
@app.post("/upload", response_model=DocumentUploadResponse)
@require_permission(Permission.CREATE_DOCUMENT)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    category: Optional[DocumentCategory] = Form(None),
    # Accept priority as a raw form field to handle both string names and numeric values
    priority: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),  # JSON string of tags
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload and process document"""
    
    # Validate file
    is_valid, error_msg = validate_file(file)
    if not is_valid:
        raise HTTPException(status_code=400, detail=error_msg)
    
    try:
        # Check for duplicate files
        temp_content = await file.read()
        await file.seek(0)  # Reset file pointer
        temp_hash = hashlib.sha256(temp_content).hexdigest()
        
        existing_doc = db.query(Document).filter_by(file_hash=temp_hash).first()
        if existing_doc:
            raise HTTPException(
                status_code=400, 
                detail=f"File already exists: {existing_doc.original_filename}"
            )
        
        # Coerce priority to enum (default to MEDIUM when not provided or invalid)
        try:
            priority_enum: Priority = Priority.MEDIUM
            if priority is not None:
                # Try numeric mapping first
                p_val = None
                if isinstance(priority, (int, float)):
                    p_val = int(priority)
                else:
                    s = str(priority).strip().lower()
                    if s.isdigit():
                        p_val = int(s)
                    else:
                        name_map = {
                            'low': Priority.LOW,
                            'medium': Priority.MEDIUM,
                            'high': Priority.HIGH,
                            'critical': Priority.CRITICAL,
                        }
                        if s in name_map:
                            priority_enum = name_map[s]
                if p_val is not None:
                    if p_val in [1, 2, 3, 4]:
                        priority_enum = Priority(p_val)
        except Exception:
            priority_enum = Priority.MEDIUM

        # Save file
        file_path, file_hash = await save_uploaded_file(file, str(current_user.id))
        
        # Parse tags
        tag_list = []
        if tags:
            try:
                tag_list = json.loads(tags)
            except:
                tag_list = [tag.strip() for tag in tags.split(',')]
        
        # Get file info
        file_size = os.path.getsize(file_path)
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # Auto-determine category if not provided
        if not category:
            category = determine_category(file.filename)
        
        # Create document record
        document = Document(
            filename=Path(file_path).name,
            original_filename=file.filename,
            file_path=file_path,
            file_size=file_size,
            mime_type=mime_type or 'application/octet-stream',
            file_hash=file_hash,
            title=title or file.filename,
            description=description,
            category=category,
            priority=priority_enum,
            tags=tag_list,
            uploaded_by=current_user.id,
            status=DocumentStatus.PROCESSING
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Create audit log
        audit_log = AuditLog(
            user_id=current_user.id,
            document_id=document.id,
            action="document_uploaded",
            entity_type="document",
            entity_id=str(document.id),
            details={
                "filename": file.filename,
                "file_size": file_size,
                "category": category.value if hasattr(category, 'value') else (str(category) if category else "unknown"),
                "priority": priority_enum.value
            }
        )
        db.add(audit_log)
        db.commit()
        
        # Process document in background
        background_tasks.add_task(process_document_content, document, db)
        
        # Create associated task for document processing
        try:
            # Import task service models
            import requests
            from config import service_config
            
            # Create task for document processing
            task_data = {
                "title": f"Process Document: {file.filename}",
                "description": f"Automated processing task for uploaded document '{file.filename}'. Includes text extraction, classification, and analysis.",
                "document_id": str(document.id),
                "assigned_to": str(current_user.id),  # Assign to uploader initially
                "priority": priority_enum.value,
                "category": category.value if category else "Document Processing",
                "task_type": "DOCUMENT_PROCESS",
                "tags": tag_list + ["auto-generated", "document-processing"],
                "task_metadata": {
                    "auto_generated": True,
                    "document_filename": file.filename,
                    "file_size": file_size,
                    "mime_type": mime_type or 'application/octet-stream'
                }
            }
            
            # Call task service to create task
            task_response = requests.post(
                f"http://localhost:{service_config.get('task_service', {}).get('port', 8004)}/tasks",
                json=task_data,
                headers={"Authorization": f"Bearer {current_user.username}"}  # Pass user context
            )
            
            if task_response.status_code == 200:
                task_id = task_response.json().get("id")
                logger.info(f"Created processing task {task_id} for document {document.id}")
            else:
                logger.warning(f"Failed to create task for document {document.id}: {task_response.text}")
                
        except Exception as task_error:
            logger.error(f"Error creating task for document {document.id}: {task_error}")
            # Don't fail the upload if task creation fails
        
        logger.info(f"Document uploaded: {file.filename} by {current_user.username}")
        
        return DocumentUploadResponse(
            document_id=str(document.id),
            filename=file.filename,
            file_size=file_size,
            status="processing",
            message="Document uploaded successfully and is being processed"
        )
        
    except Exception as e:
        # Log full stack trace for easier debugging
        import traceback
        logger.error("Error uploading document: %s\n%s", e, traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Upload failed: {e.__class__.__name__}: {str(e)}")

@app.get("/documents", response_model=List[DocumentInfo])
@require_permission(Permission.READ_DOCUMENT)
async def list_documents(
    skip: int = 0,
    limit: int = 100,
    category: Optional[DocumentCategory] = None,
    status: Optional[DocumentStatus] = None,
    search: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List documents with optional filtering"""
    
    # Base query
    query = db.query(Document)
    
    # Apply filters
    if category:
        query = query.filter(Document.category == category)
    if status:
        query = query.filter(Document.status == status)
    if search:
        query = query.filter(
            (Document.title.contains(search)) |
            (Document.original_filename.contains(search)) |
            (Document.description.contains(search))
        )
    
    # Order by creation date, newest first
    query = query.order_by(Document.created_at.desc())
    
    # Apply pagination
    documents = query.offset(skip).limit(limit).all()
    
    # Convert to response format
    result = []
    for document in documents:
        # Check access permission
        if check_document_access(document, current_user, db):
            uploader = db.query(User).filter_by(id=document.uploaded_by).first()
            result.append(DocumentInfo(
                id=str(document.id),
                filename=document.filename,
                original_filename=document.original_filename,
                file_size=document.file_size,
                mime_type=document.mime_type,
                title=document.title,
                description=document.description,
                category=document.category.value,
                priority=document.priority,
                status=document.status.value,
                uploaded_by=f"{uploader.first_name} {uploader.last_name}" if uploader else "Unknown",
                created_at=document.created_at,
                processed_at=document.processed_at,
                extracted_text=document.extracted_text[:1000] if document.extracted_text else None,
                summary=document.summary,
                language_detected=document.language_detected,
                tags=document.tags or []
            ))
    
    return result
@app.get("/documents/shared", response_model=SharedDocumentsList)
@require_permission(Permission.READ_DOCUMENT)
async def list_shared_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List documents shared by and with the current user"""
    
    # Get documents shared by the user
    shared_by_me = db.query(SharedDocument).join(
        Document, SharedDocument.document_id == Document.id
    ).filter(
        SharedDocument.shared_by == current_user.id,
        (SharedDocument.expires_at.is_(None) | (SharedDocument.expires_at > datetime.now(timezone.utc)))
    ).all()
    
    # Get documents shared with the user (directly or via department)
    shared_with_me = db.query(SharedDocument).join(
        Document, SharedDocument.document_id == Document.id
    ).filter(
        (
            (SharedDocument.shared_with_user == current_user.id) |
            (SharedDocument.shared_with_department == current_user.department)
        ),
        (SharedDocument.expires_at.is_(None) | (SharedDocument.expires_at > datetime.now(timezone.utc)))
    ).all()
    
    # Convert to response model
    shared_by_me_info = []
    shared_with_me_info = []
    
    for share in shared_by_me:
        doc = share.document
        shared_with_name = None
        if share.shared_with_user:
            user = db.query(User).filter_by(id=share.shared_with_user).first()
            if user:
                shared_with_name = f"{user.first_name} {user.last_name}"
                
        shared_by_me_info.append(SharedDocumentInfo(
            id=str(share.id),
            document_id=str(doc.id),
            document_title=doc.title,
            document_filename=doc.filename,
            shared_by=str(share.shared_by),
            shared_by_name=f"{current_user.first_name} {current_user.last_name}",
            shared_with_user=str(share.shared_with_user) if share.shared_with_user else None,
            shared_with_user_name=shared_with_name,
            shared_with_department=share.shared_with_department,
            can_edit=share.can_edit,
            created_at=share.created_at,
            expires_at=share.expires_at
        ))
    
    for share in shared_with_me:
        doc = share.document
        shared_by_user = db.query(User).filter_by(id=share.shared_by).first()
        shared_by_name = f"{shared_by_user.first_name} {shared_by_user.last_name}" if shared_by_user else "Unknown"
        
        shared_with_me_info.append(SharedDocumentInfo(
            id=str(share.id),
            document_id=str(doc.id),
            document_title=doc.title,
            document_filename=doc.filename,
            shared_by=str(share.shared_by),
            shared_by_name=shared_by_name,
            shared_with_user=str(current_user.id),
            shared_with_user_name=f"{current_user.first_name} {current_user.last_name}",
            shared_with_department=share.shared_with_department,
            can_edit=share.can_edit,
            created_at=share.created_at,
            expires_at=share.expires_at
        ))
    
    return SharedDocumentsList(
        shared_by_me=shared_by_me_info,
        shared_with_me=shared_with_me_info
    )

@app.get("/documents/{document_id}", response_model=DocumentInfo)
@require_permission(Permission.READ_DOCUMENT)
async def get_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get document information"""
    
    document = db.query(Document).filter_by(id=document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check access permission
    if not check_document_access(document, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied")
    
    uploader = db.query(User).filter_by(id=document.uploaded_by).first()
    
    return DocumentInfo(
        id=str(document.id),
        filename=document.filename,
        original_filename=document.original_filename,
        file_size=document.file_size,
        mime_type=document.mime_type,
        title=document.title,
        description=document.description,
        category=document.category.value,
        priority=document.priority,
        status=document.status.value,
        uploaded_by=f"{uploader.first_name} {uploader.last_name}" if uploader else "Unknown",
        created_at=document.created_at,
        processed_at=document.processed_at,
        extracted_text=document.extracted_text[:1000] if document.extracted_text else None,  # Limit for API
        summary=document.summary,
        language_detected=document.language_detected,
        tags=document.tags or []
    )

@app.get("/documents/{document_id}/download")
@require_permission(Permission.READ_DOCUMENT)
async def download_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download document file"""
    
    document = db.query(Document).filter_by(id=document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check access permission
    if not check_document_access(document, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied")
    
    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="Document file not found")
    
    # Log download
    audit_log = AuditLog(
        user_id=current_user.id,
        document_id=document.id,
        action="document_downloaded",
        entity_type="document",
        entity_id=str(document.id)
    )
    db.add(audit_log)
    db.commit()
    
    return FileResponse(
        path=document.file_path,
        filename=document.original_filename,
        media_type=document.mime_type
    )

@app.get("/documents/{document_id}/versions", response_model=DocumentVersionList)
@require_permission(Permission.READ_DOCUMENT)
async def get_document_versions(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get list of document versions"""
    
    document = db.query(Document).filter_by(id=document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
        
    if not check_document_access(document, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied")
        
    versions = []
    for version in document.versions:
        modified_by = db.query(User).filter_by(id=version.modified_by).first()
        modified_by_name = f"{modified_by.first_name} {modified_by.last_name}" if modified_by else "Unknown"
        
        versions.append(DocumentVersionInfo(
            id=str(version.id),
            document_id=str(document.id),
            version_number=version.version_number,
            file_size=version.file_size,
            modified_by=str(version.modified_by),
            modified_by_name=modified_by_name,
            changes_description=version.changes_description,
            created_at=version.created_at
        ))
    
    return DocumentVersionList(versions=versions)

@app.get("/documents/{document_id}/versions/{version_number}/download")
@require_permission(Permission.READ_DOCUMENT)
async def download_document_version(
    document_id: str,
    version_number: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download specific version of a document"""
    
    document = db.query(Document).filter_by(id=document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
        
    if not check_document_access(document, current_user, db):
        raise HTTPException(status_code=403, detail="Access denied")
    
    version = db.query(DocumentVersion).filter_by(
        document_id=document_id,
        version_number=version_number
    ).first()
    
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    
    if not os.path.exists(version.file_path):
        raise HTTPException(status_code=404, detail="Version file not found")
    
    # Log download
    audit_log = AuditLog(
        user_id=current_user.id,
        document_id=document.id,
        action="version_downloaded",
        entity_type="document",
        entity_id=str(document.id),
        details={"version": version_number}
    )
    db.add(audit_log)
    db.commit()
    
    return FileResponse(
        path=version.file_path,
        filename=f"{document.original_filename}.v{version_number}",
        media_type=document.mime_type
    )

@app.put("/documents/{document_id}", response_model=DocumentInfo)
async def update_document(
    document_id: str,
    document_update: DocumentUpdate,
    new_file: Optional[UploadFile] = File(None),
    changes_description: Optional[str] = Form(None),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update document information and optionally create new version"""
    
    document = db.query(Document).filter_by(id=document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check if user has edit permission
    if not check_document_access(document, current_user, db, require_edit=True):
        raise HTTPException(status_code=403, detail="Access denied - Edit permission required")
    
    # If a new file is provided, create a new version
    if new_file:
        # Validate file
        is_valid, error_msg = validate_file(new_file)
        if not is_valid:
            raise HTTPException(status_code=400, detail=error_msg)
            
        # Get next version number
        next_version = 1
        latest_version = db.query(DocumentVersion).filter_by(
            document_id=document_id
        ).order_by(DocumentVersion.version_number.desc()).first()
        
        if latest_version:
            next_version = latest_version.version_number + 1
            
        # Create version directory
        version_dir = f"{str(document_id)}/versions"
        
        # Save new version file
        file_path, file_hash = await save_uploaded_file(new_file, str(current_user.id), version_dir)
        
        # Create version record
        version = DocumentVersion(
            document_id=document_id,
            version_number=next_version,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_hash=file_hash,
            modified_by=current_user.id,
            changes_description=changes_description
        )
        db.add(version)
        
        # Update document to point to new file
        document.file_path = file_path
        document.file_size = os.path.getsize(file_path)
        document.file_hash = file_hash
        
    # Update fields
    update_data = document_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(document, field, value)
    
    db.commit()
    db.refresh(document)
    
    # Log update
    audit_log = AuditLog(
        user_id=current_user.id,
        document_id=document.id,
        action="document_updated",
        entity_type="document",
        entity_id=str(document.id),
        details={"updated_fields": list(update_data.keys())}
    )
    db.add(audit_log)
    db.commit()
    
    return document

@app.post("/documents/{document_id}/share", response_model=SharedDocumentInfo)
@require_permission(Permission.SHARE_DOCUMENT)
async def share_document(
    document_id: str,
    share_request: ShareDocumentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Share a document with a user or department"""
    
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Validate target user if sharing with user
    target_user = None
    if share_request.shared_with_user:
        target_user = db.query(User).filter(User.id == share_request.shared_with_user).first()
        if not target_user:
            raise HTTPException(status_code=404, detail="Target user not found")
            
    # Check if share already exists
    existing_share = db.query(SharedDocument).filter(
        SharedDocument.document_id == document_id,
        (
            (SharedDocument.shared_with_user == share_request.shared_with_user)
            if share_request.shared_with_user
            else (SharedDocument.shared_with_department == share_request.shared_with_department)
        )
    ).first()
    
    if existing_share:
        raise HTTPException(status_code=400, detail="Document is already shared with this user/department")
    
    try:
        # Create share record
        share = SharedDocument(
            document_id=document_id,
            shared_by=current_user.id,
            shared_with_user=share_request.shared_with_user,
            shared_with_department=share_request.shared_with_department,
            can_edit=share_request.can_edit,
            expires_at=share_request.expires_at
        )
        db.add(share)
        
        # Log sharing
        audit_log = AuditLog(
            user_id=current_user.id,
            document_id=document.id,
            action="share_document",
            entity_type="document",
            entity_id=str(document.id),
            details={
                "shared_with_user": str(share_request.shared_with_user) if share_request.shared_with_user else None,
                "shared_with_department": share_request.shared_with_department,
                "can_edit": share_request.can_edit
            }
        )
        db.add(audit_log)
        
        # Create notification for target user if sharing with specific user
        if target_user:
            notification = Notification(
                user_id=target_user.id,
                document_id=document.id,
                type=NotificationType.DOCUMENT_SHARED,
                title="Document shared with you",
                message=f"{current_user.username} has shared document '{document.filename}' with you",
                metadata={
                    "document_id": str(document.id),
                    "shared_by": str(current_user.id),
                    "can_edit": share_request.can_edit
                }
            )
            db.add(notification)
        
        db.commit()
        return SharedDocumentInfo.from_orm(share)
        
    except Exception as e:
        db.rollback()
        logger.error(f"Error sharing document: {e}")
        raise HTTPException(status_code=500, detail="Failed to share document")


@app.delete("/documents/{document_id}/share/{share_id}")
@require_permission(Permission.SHARE_DOCUMENT)
async def remove_document_share(
    document_id: str,
    share_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Remove document share"""
    
    share = db.query(SharedDocument).filter(
        SharedDocument.id == share_id,
        SharedDocument.document_id == document_id
    ).first()
    
    if not share:
        raise HTTPException(status_code=404, detail="Share not found")
    
    # Only allow removing shares if user is document owner or admin
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if document.uploaded_by != current_user.id and current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="Not authorized to remove this share")
    
    try:
        # Log share removal
        audit_log = AuditLog(
            user_id=current_user.id,
            document_id=document.id,
            action="remove_share",
            entity_type="document",
            entity_id=str(document.id),
            details={
                "share_id": str(share.id),
                "shared_with_user": str(share.shared_with_user) if share.shared_with_user else None,
                "shared_with_department": share.shared_with_department
            }
        )
        db.add(audit_log)
        
        # Remove share
        db.delete(share)
        db.commit()
        
        return {"message": "Share removed successfully"}
        
    except Exception as e:
        db.rollback()
        logger.error(f"Error removing share: {e}")
        raise HTTPException(status_code=500, detail="Failed to remove share")
        details=update_data
    db.add(audit_log)
    db.commit()
    
    logger.info(f"Document {document_id} updated by {current_user.username}")
    
    # Return updated document info
    return await get_document(document_id, current_user, db)

@app.delete("/documents/{document_id}")
@require_permission(Permission.DELETE_DOCUMENT)
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete document"""
    
    document = db.query(Document).filter_by(id=document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Check if user has permission to delete
    # Only document owner or admin can delete, regardless of sharing
    if document.uploaded_by != current_user.id and current_user.role != UserRole.ADMIN:
        raise HTTPException(status_code=403, detail="Access denied - Only document owner or admin can delete")
    
    # Delete file from disk
    try:
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
    except Exception as e:
        logger.warning(f"Could not delete file {document.file_path}: {e}")
    
    # Delete from database
    db.delete(document)
    
    # Log deletion
    audit_log = AuditLog(
        user_id=current_user.id,
        action="document_deleted",
        entity_type="document",
        entity_id=str(document.id),
        details={"filename": document.original_filename}
    )
    db.add(audit_log)
    db.commit()
    
    logger.info(f"Document {document_id} deleted by {current_user.username}")
    
    return {"message": "Document deleted successfully"}

@app.post("/search", response_model=List[DocumentInfo])
@require_permission(Permission.READ_DOCUMENT)
async def search_documents(
    search_params: DocumentSearch,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Search documents with sharing filters"""
    
    # Start with documents the user has direct access to
    base_query = db.query(Document).outerjoin(
        SharedDocument,
        Document.id == SharedDocument.document_id
    )
    
    # Build sharing filters
    sharing_conditions = []
    
    # Documents owned by user
    sharing_conditions.append(Document.uploaded_by == current_user.id)
    
    if search_params.include_shared:
        if search_params.shared_with_me:
            # Documents shared directly with user
            sharing_conditions.append(
                SharedDocument.shared_with_user == current_user.id
            )
            
        if search_params.shared_with_department:
            # Documents shared with user's department
            sharing_conditions.append(
                SharedDocument.shared_with_department == current_user.department
            )
            
        if search_params.shared_by_me:
            # Documents shared by user
            sharing_conditions.append(
                SharedDocument.shared_by == current_user.id
            )
            
    base_query = base_query.filter(or_(*sharing_conditions))
    
    # Filter expired shares
    base_query = base_query.filter(or_(
        SharedDocument.expires_at.is_(None),
        SharedDocument.expires_at > datetime.now(timezone.utc)
    ))
    
    # Admin can see all documents
    if current_user.role == UserRole.ADMIN:
        base_query = db.query(Document)
    
    # Apply search filters
    # Apply regular search filters
    if search_params.query:
        search_term = f"%{search_params.query}%"
        base_query = base_query.filter(or_(
            Document.title.ilike(search_term),
            Document.description.ilike(search_term),
            Document.original_filename.ilike(search_term),
            Document.extracted_text.ilike(search_term)
        ))
    
    if search_params.category:
        base_query = base_query.filter(Document.category == search_params.category)
    
    if search_params.priority:
        base_query = base_query.filter(Document.priority == search_params.priority)
    
    if search_params.status:
        base_query = base_query.filter(Document.status == search_params.status)
    
    if search_params.uploaded_by:
        base_query = base_query.filter(Document.uploaded_by == search_params.uploaded_by)
    
    if search_params.date_from:
        base_query = base_query.filter(Document.created_at >= search_params.date_from)
    
    if search_params.date_to:
        base_query = base_query.filter(Document.created_at <= search_params.date_to)
        
    # Order by creation date (newest first) and apply pagination
    base_query = base_query.order_by(Document.created_at.desc())
    
    # Apply pagination
    base_query = base_query.offset(search_params.offset).limit(search_params.limit)
    
    # Get unique results and convert to response model
    results = []
    seen_docs = set()
    
    for doc in base_query.all():
        if doc.id not in seen_docs:
            seen_docs.add(doc.id)
            uploader = db.query(User).filter_by(id=doc.uploaded_by).first()
            
            results.append(DocumentInfo(
                id=str(doc.id),
                filename=doc.filename,
                original_filename=doc.original_filename,
                file_size=doc.file_size,
                mime_type=doc.mime_type,
                title=doc.title,
                description=doc.description,
                category=doc.category.value,
                priority=doc.priority,
                status=doc.status.value,
                uploaded_by=f"{uploader.first_name} {uploader.last_name}" if uploader else "Unknown",
                created_at=doc.created_at,
                processed_at=doc.processed_at,
                extracted_text=doc.extracted_text[:1000] if doc.extracted_text else None,
                summary=doc.summary,
                language_detected=doc.language_detected,
                tags=doc.tags or []
            ))
    
    return results

@app.get("/categories")
async def get_document_categories():
    """Get available document categories"""
    return [{"value": cat.value, "label": cat.value.title()} for cat in DocumentCategory]

@app.get("/priorities")
async def get_document_priorities():
    """Get available document priorities"""
    return [{"value": prio.value, "label": prio.name} for prio in Priority]

@app.get("/stats")
async def get_document_stats(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get document statistics"""
    
    base_query = db.query(Document)
    
    # Filter based on user role
    if current_user.role.value != 'admin':
        base_query = base_query.filter(Document.uploaded_by == current_user.id)
    
    total_documents = base_query.count()
    
    # Status distribution
    status_stats = {}
    for status in DocumentStatus:
        count = base_query.filter(Document.status == status).count()
        status_stats[status.value] = count
    
    # Category distribution
    category_stats = {}
    for category in DocumentCategory:
        count = base_query.filter(Document.category == category).count()
        category_stats[category.value] = count
    
    # Priority distribution
    priority_stats = {}
    for priority in Priority:
        count = base_query.filter(Document.priority == priority).count()
        priority_stats[priority.name.lower()] = count
    
    return {
        "total_documents": total_documents,
        "status_distribution": status_stats,
        "category_distribution": category_stats,
        "priority_distribution": priority_stats
    }

@app.get("/health")
async def health_check(db: Session = Depends(get_db)):
    """Health check endpoint"""
    
    try:
        # Check database
        db.execute(text("SELECT 1"))
        
        # Check upload directory
        upload_dir_exists = os.path.exists(app_config.upload_directory)
        
        return {
            "status": "healthy",
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "database": "healthy",
            "upload_directory": "healthy" if upload_dir_exists else "unhealthy",
            "version": "1.0.0"
        }
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return {
            "status": "unhealthy",
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "error": str(e)
        }

if __name__ == "__main__":
    import uvicorn
    
    logger.info(f"Starting Document Service on port {service_config.document_service_port}")
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=service_config.document_service_port,
        log_level="info"
    )